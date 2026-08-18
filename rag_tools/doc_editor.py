#!/usr/bin/env python3
"""
Juvelle RAG Tool: doc_editor.py
===============================
Direct Microsoft Word (.docx) Knowledge Base editor and builder.
Allows structured editing of the master knowledge document with automatic
vector synchronization.

Usage:
    python rag_tools/doc_editor.py --view
    python rag_tools/doc_editor.py --rebuild --auto-sync
    python rag_tools/doc_editor.py --add-bullet "5. Return Policy" "Exchange Window: Exchanges accepted within 7 days of delivery." --auto-sync
"""

import os
import sys
import argparse
import logging
from typing import List, Dict, Any

# Ensure UTF-8 output encoding on Windows consoles
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

# Ensure project root is on sys.path
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from rag_tools.quick_sync import quick_sync

logger = logging.getLogger("RAGDocEditor")

DEFAULT_DOCX_PATH = os.path.join(PROJECT_ROOT, "data", "Juvelle_Knowledge_Base.docx")

def view_docx_content(file_path: str = DEFAULT_DOCX_PATH):
    """Prints a structured view of sections and bullets in the DOCX file."""
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        return

    from docx import Document
    doc = Document(file_path)
    print("\n" + "=" * 70)
    print(f"📖 JUVELLE KNOWLEDGE BASE INSPECTOR: {os.path.basename(file_path)}")
    print("=" * 70)

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        if p.style.name.startswith("Heading") or (len(text) < 60 and text[0].isdigit() and "." in text[:3]):
            print(f"\n📂 [SECTION] {text}")
            print("-" * 50)
        else:
            print(f"   • {text}")
    print("\n" + "=" * 70 + "\n")

def rebuild_default_docx(file_path: str = DEFAULT_DOCX_PATH):
    """Rebuilds the standard master Juvelle knowledge document."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Juvelle Boutique - Master Knowledge Base & FAQ")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(55, 151, 240)

    def add_sec(title, content_list):
        h = doc.add_paragraph()
        r = h.add_run(title)
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = RGBColor(30, 30, 30)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)

        for item in content_list:
            p = doc.add_paragraph(style='List Bullet')
            if ":" in item:
                bold_part, rest = item.split(":", 1)
                rb = p.add_run(bold_part + ":")
                rb.font.bold = True
                rb.font.size = Pt(10.5)
                rt = p.add_run(rest)
                rt.font.size = Pt(10.5)
            else:
                r = p.add_run(item)
                r.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(3)

    add_sec("1. Brand Identity & Product Catalog", [
        "Brand Overview: Juvelle is an exclusive women's fashion boutique based in Kerala, providing premium quality ethnic and daily wear tops at direct-to-consumer prices.",
        "Specialty Products: Exclusively women's Churidar tops crafted for daily wear, office wear, and college wear.",
        "Fabrics & Materials: Crafted with 100% breathable pure cotton and soft premium rayon blends, thoroughly tested for all-day comfort in tropical climates.",
        "Price Range: Standard retail prices range affordably between ₹399 and ₹899.",
        "Strict Exclusions: Juvelle does NOT sell sarees, frocks, jeans, t-shirts, western wear, kids wear, or men's clothing."
    ])

    add_sec("2. Shipping & Delivery Policies", [
        "Coverage Area: Delivery is available exclusively within Kerala. We do NOT ship to other Indian states (such as Bangalore, Chennai, Mumbai) or internationally.",
        "Courier Partner: All parcels are shipped safely via Delhivery courier service with end-to-end tracking.",
        "Dispatch Timeline: Orders are dispatched on the next business day following payment verification.",
        "Delivery Duration: Standard delivery takes 2 to 3 business days anywhere inside Kerala."
    ])

    add_sec("3. Ordering Process & Payment Methods", [
        "No Website: Juvelle operates directly through direct messaging on Instagram and WhatsApp without a separate e-commerce website.",
        "How to Order: Customers simply send a screenshot or photo of the top they like along with their required size (S, M, L, XL, XXL).",
        "Payment Options: 100% online advance payment via UPI (Google Pay, PhonePe, Paytm, BHIM) or direct Bank Transfer.",
        "No Cash on Delivery: Cash on Delivery (COD) is NOT available to ensure rapid order processing and next-day dispatch."
    ])

    add_sec("4. Customer Support, Sizes & Return Policy", [
        "Available Sizes: Available in sizes S (36), M (38), L (40), XL (42), and XXL (44). Custom size charts can be provided in chat upon request.",
        "Damage & Exchanges: Damaged items are replaced upon providing an unboxing parcel opening video.",
        "Customer Support Hours: Support is handled directly in chat 7 days a week.",
        "Language & Tone: Polite, friendly, and helpful. Supports English, natural Manglish (without heavy slang), and clean Malayalam script."
    ])

    doc.save(file_path)
    print(f"✅ Generated fresh master document at '{file_path}'.")

def add_bullet_to_section(section_title: str, bullet_text: str, file_path: str = DEFAULT_DOCX_PATH):
    """Appends a new bullet point to a specified or new section."""
    from docx import Document
    from docx.shared import Pt

    if not os.path.exists(file_path):
        rebuild_default_docx(file_path)

    doc = Document(file_path)
    
    # Check if section exists
    section_found = False
    for p in doc.paragraphs:
        if section_title.lower() in p.text.lower():
            section_found = True
            break

    if not section_found:
        h = doc.add_paragraph()
        r = h.add_run(section_title)
        r.font.size = Pt(14)
        r.font.bold = True
        h.paragraph_format.space_before = Pt(12)

    p = doc.add_paragraph(style='List Bullet')
    if ":" in bullet_text:
        bold_part, rest = bullet_text.split(":", 1)
        rb = p.add_run(bold_part + ":")
        rb.font.bold = True
        rb.font.size = Pt(10.5)
        rt = p.add_run(rest)
        rt.font.size = Pt(10.5)
    else:
        r = p.add_run(bullet_text)
        r.font.size = Pt(10.5)

    doc.save(file_path)
    print(f"✅ Appended bullet to section '{section_title}' in '{file_path}'.")

def main():
    parser = argparse.ArgumentParser(description="Edit or view the Juvelle Knowledge Base DOCX document.")
    parser.add_argument("--view", action="store_true", help="View current sections and bullets.")
    parser.add_argument("--rebuild", action="store_true", help="Rebuild the default master document.")
    parser.add_argument("--add-bullet", nargs=2, metavar=("SECTION", "TEXT"), help="Add a bullet point to a section.")
    parser.add_argument("--auto-sync", action="store_true", help="Auto-trigger quick vector sync after editing.")
    args = parser.parse_args()

    if args.view:
        view_docx_content()
    elif args.rebuild:
        rebuild_default_docx()
        if args.auto_sync:
            quick_sync()
    elif args.add_bullet:
        sec, text = args.add_bullet
        add_bullet_to_section(sec, text)
        if args.auto_sync:
            quick_sync()
    else:
        view_docx_content()

if __name__ == "__main__":
    main()
