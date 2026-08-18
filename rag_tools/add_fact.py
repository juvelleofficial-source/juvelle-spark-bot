#!/usr/bin/env python3
"""
Juvelle RAG Tool: add_fact.py
=============================
Instantly inject a single fact, product update, discount, or policy directly
into the active Qdrant Cloud vector database and local memory cache in <500ms.

Usage:
    python rag_tools/add_fact.py --title "Onam Offer" --content "Special 10% discount on all cotton Churidar tops with coupon ONAM10."
    python rag_tools/add_fact.py --title "New Size" --content "XXXL (46) size is now available on select rayon designs." --sync-docx
"""

import os
import sys
import time
import uuid
import argparse
import logging
from typing import Optional

# Ensure UTF-8 output encoding on Windows consoles
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

# Ensure project root is on sys.path
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from config.settings import settings
from ingestion.batch_embedder import call_vertex_batch_embeddings
from ingestion.vector_indexer import get_qdrant_client, _LOCAL_VECTOR_STORE

logger = logging.getLogger("RAGAddFact")

def add_fact(
    title: str,
    content: str,
    doc_id: Optional[str] = None,
    sync_docx: bool = False,
    verbose: bool = True
) -> str:
    """
    Injects a single fact/rule directly into Qdrant Cloud and local memory.
    
    Args:
        title: Short title or category of the fact.
        content: The text content of the fact/rule/product info.
        doc_id: Optional custom identifier.
        sync_docx: If True, also appends this fact to data/Juvelle_Knowledge_Base.docx.
        verbose: If True, prints formatted console logs.
        
    Returns:
        str: The generated chunk_id.
    """
    start_time = time.time()
    
    if not doc_id:
        doc_id = f"FACT_{int(time.time())}"
        
    chunk_id = f"{doc_id}_chk_0"
    
    # 1. Compute Embedding
    embeddings = call_vertex_batch_embeddings([content])
    embedding = embeddings[0]
    
    record = {
        "chunk_id": chunk_id,
        "doc_id": doc_id,
        "doc_title": title,
        "source_uri": "cli://instant_injection",
        "chunk_index": 0,
        "content": f"{title}: {content}" if not content.startswith(title) else content,
        "token_count": len(content.split()),
        "embedding": embedding
    }
    
    # 2. Upsert to Qdrant Cloud
    client = get_qdrant_client()
    qdrant_synced = False
    if client:
        try:
            from qdrant_client.models import PointStruct
            point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, chunk_id))
            client.upsert(
                collection_name=settings.QDRANT_COLLECTION_NAME,
                points=[
                    PointStruct(
                        id=point_id,
                        vector=[float(x) for x in embedding],
                        payload={
                            "chunk_id": record["chunk_id"],
                            "doc_id": record["doc_id"],
                            "doc_title": record["doc_title"],
                            "source_uri": record["source_uri"],
                            "chunk_index": record["chunk_index"],
                            "content": record["content"],
                            "token_count": record["token_count"]
                        }
                    )
                ]
            )
            qdrant_synced = True
        except Exception as e:
            logger.error(f"Error upserting to Qdrant: {e}")

    # 3. Add to local vector cache
    _LOCAL_VECTOR_STORE.append(record)

    # 4. Optional DOCX Append
    docx_appended = False
    if sync_docx:
        docx_path = os.path.join(PROJECT_ROOT, "data", "Juvelle_Knowledge_Base.docx")
        try:
            from docx import Document
            from docx.shared import Pt
            if os.path.exists(docx_path):
                doc = Document(docx_path)
            else:
                doc = Document()
            
            p = doc.add_paragraph(style='List Bullet')
            r_bold = p.add_run(f"{title}: ")
            r_bold.bold = True
            r_bold.font.size = Pt(10.5)
            r_text = p.add_run(content)
            r_text.font.size = Pt(10.5)
            doc.save(docx_path)
            docx_appended = True
        except Exception as e:
            logger.warning(f"Could not append to DOCX: {e}")

    elapsed_ms = (time.time() - start_time) * 1000

    if verbose:
        print("\n" + "=" * 60)
        print("➕ JUVELLE RAG: INSTANT FACT INJECTION")
        print("=" * 60)
        print(f"📌 Title:         {title}")
        print(f"📝 Content:       {content}")
        print(f"🆔 Chunk ID:      {chunk_id}")
        print(f"☁️  Qdrant Sync:   {'✅ Success' if qdrant_synced else '⚠️ Offline / Skipped'}")
        print(f"📄 DOCX Appended: {'✅ Saved' if docx_appended else ('⏩ Skipped' if not sync_docx else '❌ Failed')}")
        print(f"⏱️  Execution:     {elapsed_ms:.1f} ms")
        print("=" * 60 + "\n")

    return chunk_id

def main():
    parser = argparse.ArgumentParser(description="Instantly add a new fact to the Juvelle RAG system.")
    parser.add_argument("--title", type=str, required=True, help="Title or category for the fact.")
    parser.add_argument("--content", type=str, required=True, help="Full content/text of the fact.")
    parser.add_argument("--doc-id", type=str, default=None, help="Custom document ID.")
    parser.add_argument("--sync-docx", action="store_true", help="Also append to data/Juvelle_Knowledge_Base.docx.")
    parser.add_argument("--quiet", action="store_true", help="Suppress verbose output.")
    args = parser.parse_args()

    add_fact(
        title=args.title,
        content=args.content,
        doc_id=args.doc_id,
        sync_docx=args.sync_docx,
        verbose=not args.quiet
    )

if __name__ == "__main__":
    main()
